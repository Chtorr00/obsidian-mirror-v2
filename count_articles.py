import docx
import os
import re

DOCS_DIR = r"om_docs\Obsidian Mirror"

monthly_docs = [
    ("January", "January Obsidian Mirror Articles.docx"),
    ("February", "February Obsidian Mirror Articles.docx"),
    ("March", "March Obsidian Mirror Articles.docx"),
    ("July", "July Obsidian Mirror articles.docx"),
    ("August", "Obsidian Mirror August articles.docx"),
    ("September", "Obsidian Mirror September Articles.docx"),
    ("October", "Obsidian Mirror October Articles.docx"),
    ("November", "Obsidian Mirror November Articles.docx"),
    ("December", "December Obsidian Mirror articles.docx"),
]

# Also check Sovereign Stacks series
extra_docs = [
    ("Sovereign Stacks Series", "Sovereign Stacks Series Obsidian Mirror.docx"),
]

print("=" * 80)
print("ARTICLE COUNT PER MONTH (Heading 3 = article title)")
print("=" * 80)

total_articles = 0
for month, docname in monthly_docs:
    filepath = os.path.join(DOCS_DIR, docname)
    if not os.path.exists(filepath):
        print(f"  {month}: MISSING FILE")
        continue
    
    doc = docx.Document(filepath)
    titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and para.style and para.style.name == 'Heading 3':
            titles.append(text)
    
    total_articles += len(titles)
    print(f"\n  {month}: {len(titles)} articles")
    for i, t in enumerate(titles, 1):
        print(f"    {i:2d}. {t}")

print(f"\n  TOTAL: {total_articles} articles")

# Check sovereign stacks series
for label, docname in extra_docs:
    filepath = os.path.join(DOCS_DIR, docname)
    if not os.path.exists(filepath):
        print(f"\n  {label}: MISSING FILE")
        continue
    doc = docx.Document(filepath)
    titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and para.style and para.style.name == 'Heading 3':
            titles.append(text)
    print(f"\n  {label}: {len(titles)} articles")
    for i, t in enumerate(titles, 1):
        print(f"    {i:2d}. {t}")

# Count images per month
print("\n" + "=" * 80)
print("IMAGE COUNT PER MONTH")
print("=" * 80)

IMG_DIR = r"public\images"
if os.path.exists(IMG_DIR):
    images = sorted(os.listdir(IMG_DIR))
    month_map = {}
    for img in images:
        name_lower = img.lower()
        for m in ['january', 'jan', 'february', 'feb', 'march', 'mar', 'april', 'apr', 
                   'may', 'june', 'jun', 'july', 'jul', 'august', 'aug', 'september', 'sept', 'sep',
                   'october', 'oct', 'november', 'nov', 'december', 'dec']:
            if name_lower.startswith(m):
                # Normalize to month name
                month_norm = {
                    'jan': 'January', 'january': 'January',
                    'feb': 'February', 'february': 'February',
                    'mar': 'March', 'march': 'March',
                    'apr': 'April', 'april': 'April',
                    'may': 'May',
                    'jun': 'June', 'june': 'June',
                    'jul': 'July', 'july': 'July',
                    'aug': 'August', 'august': 'August',
                    'sep': 'September', 'sept': 'September', 'september': 'September',
                    'oct': 'October', 'october': 'October',
                    'nov': 'November', 'november': 'November',
                    'dec': 'December', 'december': 'December',
                }.get(m, m)
                if month_norm not in month_map:
                    month_map[month_norm] = []
                month_map[month_norm].append(img)
                break
    
    for m in ['January', 'February', 'March', 'April', 'May', 'June', 'July', 
              'August', 'September', 'October', 'November', 'December']:
        imgs = month_map.get(m, [])
        if imgs:
            print(f"  {m}: {len(imgs)} images")
    
    print(f"\n  TOTAL: {sum(len(v) for v in month_map.values())} images")
